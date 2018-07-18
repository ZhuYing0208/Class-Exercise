
# coding: utf-8

# In[2]:


from docx import Document


# In[3]:


document=Document()


# In[4]:


document.add_paragraph('Hello')


# In[5]:


document.save('test_write_docx.docx')


# In[41]:


path='C:/Users/10523/Desktop/docx.docx'


# In[42]:


document= Document(path)


# In[8]:


for para in document.paragraphs:
    print(para.text)


# In[24]:


from pyquery import PyQuery as pq


# In[44]:


html=pq('C:\Users\10523\Desktop\Python.html')


# In[31]:


html.text()


# In[27]:


html=pq(url='https://news.qq.com/a/20180522/023866.htm')


# In[28]:


html.text


# In[29]:


html('title').text()


# In[47]:


from pyquery import PyQuery as pq


# In[48]:


import codecs


# In[49]:


f=codec.open('C:\Users\10523\Desktop\Python.html','r','utf-8')

